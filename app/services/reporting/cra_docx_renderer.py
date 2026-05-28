from __future__ import annotations

from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile
from xml.sax.saxutils import escape


def render_docx(path: Path, report: dict[str, Any]) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        return _render_python_docx(path, report)
    except ModuleNotFoundError:
        return _render_minimal_docx(path, report)


def _render_python_docx(path: Path, report: dict[str, Any]) -> Path:
    from docx import Document

    doc = Document()
    summary = report["summary"]
    narrative = report["narrative"]
    doc.add_heading("Copilot Readiness Assessment", 0)
    doc.add_paragraph(f"Prepared for: {summary['customer_name']}")
    doc.add_paragraph(f"Assessment date: {summary.get('assessment_date') or '-'}")
    doc.add_paragraph("Prepared by: CRA Platform")
    doc.add_page_break()
    doc.add_paragraph("Table of Contents")
    doc.add_heading("Executive Summary", 1)
    doc.add_paragraph(narrative["executive_summary"])
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    for key, value in [
        ("Overall readiness", f"{summary['overall_readiness']}%"),
        ("Readiness status", summary["readiness_status"]),
        ("Pass / Warning / Fail", f"{summary['pass_total']} / {summary['warning_total']} / {summary['fail_total']}"),
    ]:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value
    doc.add_heading("Analytics & Charts", 1)
    for name, items in report["analytics"].items():
        doc.add_heading(name.replace("_", " ").title(), 2)
        doc.add_paragraph(str(items))
    doc.add_heading("Detailed Assessment", 1)
    for service, findings in report["sections"].items():
        doc.add_heading(service, 2)
        table = doc.add_table(rows=1, cols=4)
        for idx, header in enumerate(["Parameter", "Severity", "Finding", "Recommendation"]):
            table.rows[0].cells[idx].text = header
        for finding in findings:
            row = table.add_row().cells
            row[0].text = finding["title"]
            row[1].text = finding["severity"]
            row[2].text = finding["finding"]
            row[3].text = finding["recommendation"]
    doc.add_heading("Conclusion", 1)
    doc.add_paragraph(narrative["conclusion"])
    doc.save(path)
    return path


def _p(text: str, style: str | None = None) -> str:
    pstyle = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ""
    return f"<w:p>{pstyle}<w:r><w:t>{escape(str(text))}</w:t></w:r></w:p>"


def _render_minimal_docx(path: Path, report: dict[str, Any]) -> Path:
    summary = report["summary"]
    narrative = report["narrative"]
    body = [
        _p("Copilot Readiness Assessment", "Title"),
        _p(f"Prepared for: {summary['customer_name']}"),
        _p(f"Assessment date: {summary.get('assessment_date') or '-'}"),
        _p("Prepared by: CRA Platform"),
        _p("Table of Contents", "Heading1"),
        _p("Executive Summary", "Heading1"),
        _p(narrative["executive_summary"]),
        _p("Analytics & Charts", "Heading1"),
        _p(str(report["analytics"])),
        _p("Detailed Assessment", "Heading1"),
    ]
    for service, findings in report["sections"].items():
        body.append(_p(service, "Heading2"))
        for finding in findings:
            body.append(_p(f"{finding['title']} | {finding['severity']} | {finding['finding']} | {finding['recommendation']}"))
    body.extend([_p("Conclusion", "Heading1"), _p(narrative["conclusion"])])
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"<w:body>{''.join(body)}<w:sectPr/></w:body></w:document>"
    )
    with ZipFile(path, "w", ZIP_DEFLATED) as docx:
        docx.writestr("[Content_Types].xml", '<?xml version="1.0" encoding="UTF-8"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/></Types>')
        docx.writestr("_rels/.rels", '<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>')
        docx.writestr("word/document.xml", document_xml)
    return path
